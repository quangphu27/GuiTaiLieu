import os
import re
import time
from openai import OpenAI
from dotenv import load_dotenv

_load_dotenv_done = False
_api_key = None
_client = None

def _load_config():
    global _load_dotenv_done, _api_key, _client
    
    if _load_dotenv_done and _client is not None:
        return _api_key, _client
    
    if _load_dotenv_done and _client is None and _api_key:
        try:
            _client = OpenAI(api_key=_api_key)
        except RuntimeError as e:
            if "can't register atexit" in str(e) or "after shutdown" in str(e):
                return _api_key, None
            raise
        except Exception:
            return _api_key, None
    
    if _load_dotenv_done:
        return _api_key, _client
    
    try:
        load_dotenv()
        _api_key = os.getenv('OPENAI_API_KEY')
        if _api_key:
            try:
                _client = OpenAI(api_key=_api_key)
            except RuntimeError as e:
                if "can't register atexit" in str(e) or "after shutdown" in str(e):
                    _client = None
                else:
                    _client = None
            except Exception:
                _client = None
        _load_dotenv_done = True
    except Exception:
        _load_dotenv_done = True
    
    return _api_key, _client

def get_openai_client():
    api_key, client = _load_config()
    return client

def is_valid_text(text):
    if not text or not text.strip():
        return False
    
    text = text.strip()
    
    if len(text) < 10:
        return False
    
    letter_chars = 0
    invalid_chars = 0
    
    vietnamese_chars = 'àáảãạăằắẳẵặâầấẩẫậèéẻẽẹêềếểễệìíỉĩịòóỏõọôồốổỗộơờớởỡợùúủũụưừứửữựỳýỷỹỵđĐÀÁẢÃẠĂẰẮẲẴẶÂẦẤẨẪẬÈÉẺẼẸÊỀẾỂỄỆÌÍỈĨỊÒÓỎÕỌÔỒỐỔỖỘƠỜỚỞỠỢÙÚỦŨỤƯỪỨỬỮỰỲÝỶỸỴĐ'
    
    for char in text:
        if char.isprintable() or char in '\n\r\t':
            if char.isalpha() or char in vietnamese_chars:
                letter_chars += 1
        else:
            invalid_chars += 1
    
    if invalid_chars > len(text) * 0.1:
        return False
    
    if letter_chars < len(text) * 0.2:
        return False
    
    special_chars_pattern = r'[^a-zA-Z0-9' + re.escape(vietnamese_chars) + r'\s.,;:!?\-_=+()\[\]{}|/\\]{20,}'
    if re.search(special_chars_pattern, text):
        return False
    
    control_count = sum(1 for c in text if ord(c) < 32 and c not in '\n\r\t')
    if control_count > len(text) * 0.05:
        return False
    
    return True

def extract_text_from_file(filepath):
    try:
        if not filepath or not os.path.exists(filepath):
            return None
        
        file_ext = os.path.splitext(filepath)[1].lower()
        
        if file_ext == '.pdf':
            text = None
            
            try:
                import pdfplumber
                import warnings
                with warnings.catch_warnings():
                    warnings.simplefilter("ignore")
                    with pdfplumber.open(filepath) as pdf:
                        text_parts = []
                        for page in pdf.pages[:10]:
                            try:
                                page_text = page.extract_text()
                                if page_text and page_text.strip():
                                    text_parts.append(page_text.strip())
                                try:
                                    for table in page.extract_tables() or []:
                                        table_text = '\n'.join([' '.join([str(cell) if cell else '' for cell in row]) for row in table])
                                        if table_text.strip():
                                            text_parts.append(table_text.strip())
                                except:
                                    pass
                            except:
                                continue
                        if text_parts:
                            text = '\n'.join(text_parts)
                            if text.strip() and is_valid_text(text):
                                return text[:5000]
            except:
                pass
            
            if not text:
                try:
                    import pypdf
                    import warnings
                    with warnings.catch_warnings():
                        warnings.simplefilter("ignore")
                        with open(filepath, 'rb') as file:
                            pdf_reader = pypdf.PdfReader(file)
                            text_parts = []
                            for page in pdf_reader.pages[:10]:
                                try:
                                    page_text = page.extract_text()
                                    if page_text and page_text.strip():
                                        text_parts.append(page_text.strip())
                                except:
                                    continue
                            if text_parts:
                                text = '\n'.join(text_parts)
                                if text.strip() and is_valid_text(text):
                                    return text[:5000]
                except:
                    pass
            
            if not text:
                try:
                    import PyPDF2
                    import warnings
                    with warnings.catch_warnings():
                        warnings.simplefilter("ignore")
                        with open(filepath, 'rb') as file:
                            pdf_reader = PyPDF2.PdfReader(file)
                            text_parts = []
                            for page in pdf_reader.pages[:10]:
                                try:
                                    page_text = page.extract_text()
                                    if page_text and page_text.strip():
                                        text_parts.append(page_text.strip())
                                except:
                                    continue
                            if text_parts:
                                text = '\n'.join(text_parts)
                                if text.strip() and is_valid_text(text):
                                    return text[:5000]
                except:
                    pass
            
            return None
        
        elif file_ext == '.docx':
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(filepath)
                text_parts = []
                for para in doc.paragraphs[:100]:
                    if para.text.strip():
                        text_parts.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            text_parts.append(row_text)
                text = '\n'.join(text_parts)
                return text[:5000]
            except:
                return None
        
        elif file_ext == '.doc':
            try:
                doc_text = []
                
                try:
                    with open(filepath, 'rb') as f:
                        content = f.read()
                    
                    ascii_patterns = re.findall(rb'[\x20-\x7E]{4,}', content)
                    for pattern in ascii_patterns[:200]:
                        try:
                            text_str = pattern.decode('utf-8', errors='ignore').strip()
                            if (len(text_str) >= 3 and 
                                any(c.isalpha() for c in text_str) and 
                                len(text_str) <= 500 and
                                not text_str.replace(' ', '').isdigit()):
                                doc_text.append(text_str)
                        except:
                            continue
                    
                    try:
                        utf16_content = content.decode('utf-16le', errors='ignore')
                        lines = utf16_content.split('\n')
                        for line in lines:
                            line = line.strip()
                            if (len(line) >= 3 and 
                                any(c.isalpha() for c in line) and 
                                len(line) <= 500):
                                doc_text.append(line)
                    except:
                        pass
                    
                    try:
                        from oletools.olefile import OleFileIO
                        ole = OleFileIO(filepath)
                        for stream_name in ole.listdir():
                            try:
                                stream_data = ole.openstream(stream_name).read()
                                for encoding in ['utf-8', 'utf-16le', 'latin-1']:
                                    try:
                                        stream_text = stream_data.decode(encoding, errors='ignore')
                                        paragraphs = stream_text.split('\n')
                                        for para in paragraphs:
                                            para = para.strip()
                                            if (len(para) >= 3 and 
                                                any(c.isalpha() for c in para) and 
                                                len(para) <= 500):
                                                doc_text.append(para)
                                        break
                                    except:
                                        continue
                            except:
                                continue
                        ole.close()
                    except:
                        pass
                    
                    unique_text = []
                    seen = set()
                    for t in doc_text:
                        t_clean = t.strip()
                        t_lower = t_clean.lower()
                        if (t_lower not in seen and 
                            len(t_clean) >= 3 and 
                            len(t_clean) <= 500 and
                            any(c.isalpha() for c in t_clean) and
                            not all(c in '.,;:!?-_=+()[]{}|/\\' for c in t_clean)):
                            seen.add(t_lower)
                            unique_text.append(t_clean)
                    
                    text = '\n'.join(unique_text[:300])
                    if text.strip():
                        return text[:5000]
                    
                except:
                    pass
                
                return None
                
            except:
                return None
        
        elif file_ext == '.xlsx':
            try:
                from openpyxl import load_workbook
                wb = load_workbook(filepath, read_only=True, data_only=True)
                text_parts = []
                for sheet_name in wb.sheetnames[:3]:
                    sheet = wb[sheet_name]
                    text_parts.append(f"Sheet: {sheet_name}")
                    for row in sheet.iter_rows(max_row=50, values_only=True):
                        row_text = ' '.join([str(cell) for cell in row if cell and str(cell).strip()])
                        if row_text:
                            text_parts.append(row_text)
                wb.close()
                text = '\n'.join(text_parts)
                return text[:5000]
            except:
                return None
        
        elif file_ext == '.xls':
            try:
                import xlrd
                workbook = xlrd.open_workbook(filepath, on_demand=False, formatting_info=False)
                text_parts = []
                for sheet_idx in range(min(3, workbook.nsheets)):
                    sheet = workbook.sheet_by_index(sheet_idx)
                    text_parts.append(f"Sheet: {sheet.name}")
                    for row_idx in range(min(50, sheet.nrows)):
                        row_values = []
                        for col_idx in range(sheet.ncols):
                            cell_value = sheet.cell_value(row_idx, col_idx)
                            if isinstance(cell_value, float) and sheet.cell_type(row_idx, col_idx) == xlrd.XL_CELL_DATE:
                                try:
                                    date_tuple = xlrd.xldate_as_tuple(cell_value, workbook.datemode)
                                    cell_value = f"{date_tuple[2]}/{date_tuple[1]}/{date_tuple[0]}"
                                except:
                                    pass
                            if cell_value and str(cell_value).strip():
                                row_values.append(str(cell_value).strip())
                        if row_values:
                            text_parts.append(' '.join(row_values))
                text = '\n'.join(text_parts)
                return text[:5000]
            except:
                return None
        
        return None
    except:
        return None

def suggest_units_from_document_streaming(document, all_units):
    try:
        if not all_units:
            yield {
                'suggested_ids': [],
                'has_suggestions': False,
                'message': 'Không có đơn vị phù hợp',
                'is_fallback': False,
                'chunk_index': 0,
                'is_final': True
            }
            return
        
        document_name = document.get('name', '')
        filepath = document.get('filepath', '')
        document_content = extract_text_from_file(filepath)
        
        if not document_content:
            yield suggest_units_fallback(document_name, None, all_units)
            return
        
        units_info = "\n".join([
            f"Index {i}: {unit['name']} (Mã: {unit.get('code', 'N/A')})"
            for i, unit in enumerate(all_units)
        ])
        
        current_key, client = _load_config()
        if not client:
            yield suggest_units_fallback(document_name, document_content, all_units)
            return
        
        chunk_size = 10000
        chunks = []
        for i in range(0, len(document_content), chunk_size):
            chunks.append(document_content[i:i + chunk_size])
        
        all_suggested_ids = set()
        has_first_success = False
        
        for chunk_idx, chunk_content in enumerate(chunks):
            try:
                prompt = f"""Chỉ chọn các đơn vị liên quan từ danh sách dựa trên nội dung tài liệu sau. Nếu không có đơn vị nào liên quan, trả về NONE.
Chỉ trả về index, phân tách bằng dấu phẩy, tối đa 5 index, không giải thích gì thêm.

NỘI DUNG TÀI LIỆU (Phần {chunk_idx + 1}/{len(chunks)}):
{chunk_content}

DANH SÁCH ĐƠN VỊ:
{units_info}"""
                
                max_retries = 1
                retry_count = 0
                result = None
                
                while retry_count <= max_retries:
                    try:
                        response = client.chat.completions.create(
                            model="gpt-4o-mini",
                            messages=[
                                {"role": "system", "content": "Chỉ trả về index của các đơn vị liên quan, phân tách bằng dấu phẩy, tối đa 5 index. Nếu không có đơn vị nào liên quan, trả về NONE. Không giải thích gì thêm."},
                                {"role": "user", "content": prompt}
                            ],
                            max_tokens=150,
                            temperature=0.2
                        )
                        
                        result = response.choices[0].message.content.strip().upper()
                        break
                        
                    except Exception as e:
                        error_msg = str(e)
                        is_rate_limit = ('429' in error_msg or 'rate_limit' in error_msg.lower() or 
                                        'Rate limit' in error_msg or 'RPM' in error_msg or 'TPM' in error_msg)
                        
                        if is_rate_limit and retry_count < max_retries:
                            wait_time = 25
                            if 'try again in' in error_msg.lower() or 'Please try again in' in error_msg:
                                time_match = re.search(r'(\d+)\s*s(?:ec)?', error_msg, re.IGNORECASE)
                                if time_match:
                                    wait_time = int(time_match.group(1)) + 5
                            
                            time.sleep(wait_time)
                            retry_count += 1
                            continue
                        else:
                            if chunk_idx == 0 and not has_first_success:
                                yield suggest_units_fallback(document_name, document_content, all_units)
                                return
                            break
                
                if result and 'NONE' not in result:
                    indices = [int(idx.strip()) for idx in result.split(',') if idx.strip().isdigit() and int(idx.strip()) < len(all_units)]
                    chunk_suggested_ids = [all_units[i]['id'] for i in indices if i < len(all_units)]
                    
                    for unit_id in chunk_suggested_ids:
                        all_suggested_ids.add(unit_id)
                    
                    has_first_success = True
                    
                    final_ids = list(all_suggested_ids)[:5]
                    yield {
                        'suggested_ids': final_ids,
                        'has_suggestions': len(final_ids) > 0,
                        'message': '',
                        'is_fallback': False,
                        'chunk_index': chunk_idx + 1,
                        'total_chunks': len(chunks),
                        'is_final': chunk_idx == len(chunks) - 1
                    }
            
            except Exception:
                if chunk_idx == 0 and not has_first_success:
                    yield suggest_units_fallback(document_name, document_content, all_units)
                    return
        
        if not all_suggested_ids:
            yield {
                'suggested_ids': [],
                'has_suggestions': False,
                'message': 'Không có đơn vị phù hợp',
                'is_fallback': False,
                'chunk_index': len(chunks),
                'is_final': True
            }
    
    except Exception:
        filepath = document.get('filepath', '')
        document_content = extract_text_from_file(filepath) if filepath else None
        yield suggest_units_fallback(document.get('name', ''), document_content, all_units)

def suggest_units_from_document(document, all_units):
    try:
        if not all_units:
            return {
                'suggested_ids': [],
                'has_suggestions': False,
                'message': 'Không có đơn vị phù hợp',
                'is_fallback': False
            }
        
        document_name = document.get('name', '')
        filepath = document.get('filepath', '')
        
        document_content = extract_text_from_file(filepath)
        
        units_info = "\n".join([
            f"Index {i}: {unit['name']} (Mã: {unit.get('code', 'N/A')})"
            for i, unit in enumerate(all_units)
        ])
        
        if document_content:
            prompt = f"""Chỉ chọn các đơn vị liên quan từ danh sách dựa trên nội dung tài liệu sau. Nếu không có đơn vị nào liên quan, trả về NONE.
Chỉ trả về index, phân tách bằng dấu phẩy, tối đa 5 index, không giải thích gì thêm.

NỘI DUNG TÀI LIỆU:
{document_content[:10000]}

DANH SÁCH ĐƠN VỊ:
{units_info}"""
        else:
            prompt = f"""Chỉ chọn các đơn vị liên quan từ danh sách dựa trên tên tài liệu sau. Nếu không có đơn vị nào liên quan, trả về NONE.
Chỉ trả về index, phân tách bằng dấu phẩy, tối đa 5 index, không giải thích gì thêm.

TÊN TÀI LIỆU: "{document_name}"

DANH SÁCH ĐƠN VỊ:
{units_info}"""
        
        current_key, client = _load_config()
        
        if not client:
            return suggest_units_fallback(document.get('name', ''), document_content, all_units)
        
        result = None
        max_retries = 1
        retry_count = 0
        
        while retry_count <= max_retries:
            try:
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "Chỉ trả về index của các đơn vị liên quan, phân tách bằng dấu phẩy, tối đa 5 index. Nếu không có đơn vị nào liên quan, trả về NONE. Không giải thích gì thêm."},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=150,
                    temperature=0.2
                )
                
                result = response.choices[0].message.content.strip().upper()
                break
                
            except Exception as e:
                error_msg = str(e)
                is_rate_limit = ('429' in error_msg or 'rate_limit' in error_msg.lower() or 
                                'Rate limit' in error_msg or 'RPM' in error_msg or 'TPM' in error_msg)
                
                if is_rate_limit and retry_count < max_retries:
                    wait_time = 25
                    if 'try again in' in error_msg.lower() or 'Please try again in' in error_msg:
                        time_match = re.search(r'(\d+)\s*s(?:ec)?', error_msg, re.IGNORECASE)
                        if time_match:
                            wait_time = int(time_match.group(1)) + 5
                    
                    time.sleep(wait_time)
                    retry_count += 1
                    continue
                else:
                    return suggest_units_fallback(document.get('name', ''), document_content, all_units)
        
        if not result or 'NONE' in result:
            return {
                'suggested_ids': [],
                'has_suggestions': False,
                'message': 'Không có đơn vị phù hợp',
                'is_fallback': False
            }
        
        indices = [int(idx.strip()) for idx in result.split(',') if idx.strip().isdigit() and int(idx.strip()) < len(all_units)]
        
        if not indices:
            return {
                'suggested_ids': [],
                'has_suggestions': False,
                'message': 'Không có đơn vị phù hợp',
                'is_fallback': False
            }
        
        suggested_ids = [all_units[i]['id'] for i in indices if i < len(all_units)][:5]
        
        return {
            'suggested_ids': suggested_ids,
            'has_suggestions': True,
            'message': '',
            'is_fallback': False
        }
    
    except:
        filepath = document.get('filepath', '')
        document_content = extract_text_from_file(filepath) if filepath else None
        return suggest_units_fallback(document.get('name', ''), document_content, all_units)

def suggest_units_fallback(document_name, document_content, all_units):
    if not all_units:
        return {
            'suggested_ids': [],
            'has_suggestions': False,
            'message': 'Không có đơn vị phù hợp',
            'is_fallback': True
        }
    
    doc_name_lower = (document_name or '').lower()
    doc_content_lower = (document_content or '').lower()
    
    suggested_ids = []
    
    keywords = ['ngân hàng', 'tài chính', 'kinh doanh', 'hành chính', 'đào tạo', 'tồn kho', 
                'kế toán', 'nhân sự', 'lương', 'hợp đồng', 'lao động', 'quy định', 'dự án']
    
    search_text = doc_content_lower if document_content else doc_name_lower
    
    for keyword in keywords:
        if keyword in search_text:
            for unit in all_units:
                unit_name_lower = (unit.get('name', '') or '').lower()
                unit_code_lower = (unit.get('code', '') or '').lower()
                
                if keyword in unit_name_lower or keyword in unit_code_lower:
                    if unit['id'] not in suggested_ids:
                        suggested_ids.append(unit['id'])
    
    if not suggested_ids:
        return {
            'suggested_ids': [],
            'has_suggestions': False,
            'message': 'Không có đơn vị phù hợp',
            'is_fallback': True
        }
    
    return {
        'suggested_ids': suggested_ids[:5],
        'has_suggestions': True,
        'message': '',
        'is_fallback': True
    }
